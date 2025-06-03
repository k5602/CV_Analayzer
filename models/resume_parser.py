import os
import re
try:
    import spacy
    SPACY_AVAILABLE = True
except ImportError:
    SPACY_AVAILABLE = False
import PyPDF2
import pdfplumber
import docx
from typing import Dict, List, Union, Optional, Tuple
from loguru import logger
try:
    import pytesseract
    from PIL import Image
    PYTESSERACT_AVAILABLE = True
except ImportError:
    PYTESSERACT_AVAILABLE = False
import io
import pandas as pd
import fitz  # PyMuPDF for PDF processing
import tempfile
import shutil
from pathlib import Path

class ResumeParser:
    """
    A class for parsing resumes in various formats (PDF, DOCX, TXT) and extracting
    structured information such as contact details, skills, work experience, and education.
    """

    def __init__(self):
        """Initialize the ResumeParser with NLP models and skills database."""
        self.nlp = None
        self.skills_db = None
        self.use_ocr = True  # Default to using OCR for image-based PDFs

        # Regex patterns for various data extraction
        self._compile_regex_patterns()

        try:
            # Load skills database (can be expanded)
            logger.info("Loading skills database...")
            self._load_skills_database()

            # Load spaCy NLP model for entity recognition if available
            if SPACY_AVAILABLE:
                try:
                    logger.info("Loading spaCy NLP model...")
                    self.nlp = spacy.load("en_core_web_sm")
                    logger.success("SpaCy model loaded successfully")
                except Exception as e:
                    logger.warning(f"Could not load spaCy model: {e}. Some functionality will be limited.")
            else:
                logger.warning("SpaCy not available. Some functionality will be limited.")
                
            # Check if OCR is available
            if not PYTESSERACT_AVAILABLE:
                logger.warning("Pytesseract not available. OCR functionality will be limited.")
                self.use_ocr = False

            logger.success("ResumeParser initialized successfully")
        except Exception as e:
            logger.error(f"Error initializing ResumeParser: {e}")
            # Continue with limited functionality rather than failing completely
            logger.warning("Continuing with limited functionality")

    def _load_skills_database(self):
        """Load or create a database of skills for matching."""
        # This could be expanded to load from a file or external source
        self.skills_db = set([
            # Programming languages
            "python", "java", "javascript", "c++", "c#", "ruby", "php", "swift", "kotlin", "go",
            "typescript", "scala", "r", "perl", "rust", "dart", "html", "css", "sql",

            # Frameworks and libraries
            "react", "angular", "vue", "django", "flask", "spring", "node.js", "express",
            "tensorflow", "pytorch", "pandas", "numpy", "scikit-learn", "keras", ".net", "laravel",

            # Cloud & DevOps
            "aws", "azure", "google cloud", "docker", "kubernetes", "jenkins", "terraform", "ansible",
            "ci/cd", "devops", "microservices", "serverless",

            # Databases
            "mysql", "postgresql", "mongodb", "oracle", "sql server", "sqlite", "redis", "cassandra",
            "dynamodb", "firebase",

            # Tools & Software
            "git", "github", "gitlab", "bitbucket", "jira", "confluence", "slack", "trello",
            "visual studio code", "intellij", "eclipse", "photoshop", "illustrator", "figma",

            # Soft skills
            "leadership", "communication", "teamwork", "problem-solving", "critical thinking",
            "time management", "creativity", "adaptability", "project management", "agile",
            "scrum", "kanban",

            # Data science/AI
            "machine learning", "deep learning", "artificial intelligence", "data science",
            "data analysis", "natural language processing", "computer vision", "big data",
            "data engineering", "data visualization", "statistics", "analytics",

            # Mobile development
            "ios", "android", "react native", "flutter", "mobile development", "app development",
            "cross-platform", "pwa",

            # Other technical skills
            "rest api", "graphql", "oauth", "jwt", "web services", "soa", "mvc", "orm",
            "responsive design", "web accessibility", "ui/ux", "frontend", "backend",
            "full-stack", "testing", "qa", "security", "blockchain"
        ])

    def _compile_regex_patterns(self):
        """Compile regex patterns for extracting various resume elements."""
        # Contact information patterns
        self.email_pattern = re.compile(r'[\w\.-]+@[\w\.-]+\.\w+')
        self.phone_pattern = re.compile(r'(\+\d{1,3}[-\.\s]?)?(\(?\d{3}\)?[-\.\s]?\d{3}[-\.\s]?\d{4}|\d{10})')
        self.linkedin_pattern = re.compile(r'linkedin\.com/in/[\w-]+')
        self.github_pattern = re.compile(r'github\.com/[\w-]+')

        # Section header patterns
        self.education_headers = re.compile(r'education|academic|qualification', re.IGNORECASE)
        self.experience_headers = re.compile(r'experience|employment|work history|career|professional', re.IGNORECASE)
        self.skills_headers = re.compile(r'skills|expertise|proficiency|competency|technical', re.IGNORECASE)
        self.projects_headers = re.compile(r'projects|portfolio|works', re.IGNORECASE)
        self.summary_headers = re.compile(r'summary|objective|profile|about', re.IGNORECASE)

        # Date patterns
        self.date_pattern = re.compile(r'(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|January|February|March|April|May|June|July|August|September|October|November|December)[a-z]*\.?\s*\d{4}\s*(-|–|to)\s*(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|January|February|March|April|May|June|July|August|September|October|November|December)[a-z]*\.?\s*\d{4}|\d{1,2}/\d{4}\s*(-|–|to)\s*\d{1,2}/\d{4}|\d{4}\s*(-|–|to)\s*\d{4}|\d{4}\s*(-|–|to)\s*(present|current|now)', re.IGNORECASE)

    def parse_resume(self, file_path: str) -> Dict:
        """
        Parse a resume file and extract structured information.

        Args:
            file_path: Path to the resume file (PDF, DOCX, or TXT)

        Returns:
            Dictionary containing extracted resume information
        """
        try:
            if not os.path.exists(file_path):
                logger.error(f"File not found: {file_path}")
                return {"error": "File not found"}

            # Extract text based on file extension
            file_ext = os.path.splitext(file_path)[1].lower()
            if file_ext == '.pdf':
                text = self._extract_text_from_pdf(file_path)
                has_problematic_elements = self._check_pdf_for_ats_issues(file_path)
            elif file_ext == '.docx':
                text = self._extract_text_from_docx(file_path)
                has_problematic_elements = False  # Need implementation for DOCX checking
            elif file_ext == '.txt':
                text = self._extract_text_from_txt(file_path)
                has_problematic_elements = False
            else:
                logger.error(f"Unsupported file format: {file_ext}")
                return {"error": f"Unsupported file format: {file_ext}"}

            if not text.strip():
                logger.error("No text could be extracted from the resume")
                return {"error": "No text could be extracted from the resume"}

            # Process extracted text to get structured information
            result = self._process_text(text)

            # Add file format information
            result["file_format"] = {
                "extension": file_ext,
                "has_problematic_elements": has_problematic_elements
            }

            return result

        except Exception as e:
            logger.error(f"Error parsing resume: {e}")
            return {"error": f"Error parsing resume: {str(e)}"}

    def _extract_text_from_pdf(self, pdf_path: str) -> str:
        """Extract text content from a PDF file."""
        text = ""

        # Try pdfplumber first (better for formatted text)
        try:
            with pdfplumber.open(pdf_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text() or ""
                    text += page_text + "\n"
        except Exception as e:
            logger.warning(f"pdfplumber extraction failed, trying PyPDF2: {e}")

            # Fallback to PyPDF2
            try:
                with open(pdf_path, "rb") as file:
                    reader = PyPDF2.PdfReader(file)
                    for page_num in range(len(reader.pages)):
                        page = reader.pages[page_num]
                        text += page.extract_text() + "\n"
            except Exception as e2:
                logger.error(f"PyPDF2 extraction also failed: {e2}")

        # If we got less than 100 characters or suspicious content, the PDF might be image-based
        # Check if text is suspiciously short or mostly whitespace/special chars
        if len(text.strip()) < 100 or self._appears_to_be_image_pdf(text):
            logger.info("PDF appears to be image-based, attempting OCR extraction")
            ocr_text = self._extract_text_using_ocr(pdf_path)
            if ocr_text:
                text = ocr_text

        return text

    def _check_pdf_for_ats_issues(self, pdf_path: str) -> bool:
        """
        Check if PDF has elements that might cause ATS issues like images,
        tables, non-standard fonts, etc.
        """
        try:
            # First check if this is an image-based PDF
            if self._is_image_based_pdf(pdf_path):
                logger.info("PDF is image-based (scanned document)")
                return True
                
            with pdfplumber.open(pdf_path) as pdf:
                for page in pdf.pages:
                    # Check for tables
                    try:
                        if page.find_tables():
                            return True
                    except:
                        pass
                    
                    # Check for images (approximation)
                    try:
                        if page.images:
                            return True
                    except:
                        pass
                    
                    # Complex layouts could be detected by analyzing the bounding boxes
                    # This is a simplified implementation
                    try:
                        if len(page.chars) > 0:
                            x_positions = set([char['x0'] for char in page.chars])
                            # If there are many different x positions, it might indicate a complex layout
                            if len(x_positions) > 15:  # arbitrary threshold
                                return True
                    except:
                        pass
            
            return False
        except Exception as e:
            logger.error(f"Error checking PDF for ATS issues: {e}")
            return True  # Assume issues if we can't check properly

    def _extract_text_from_docx(self, docx_path: str) -> str:
        """Extract text content from a DOCX file."""
        text = ""
        try:
            doc = docx.Document(docx_path)
            for para in doc.paragraphs:
                text += para.text + "\n"

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {e}")

        return text

    def _extract_text_from_txt(self, txt_path: str) -> str:
        """Extract text content from a plain text file."""
        try:
            with open(txt_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            # Try different encodings if UTF-8 fails
            encodings = ['latin-1', 'iso-8859-1', 'cp1252']
            for encoding in encodings:
                try:
                    with open(txt_path, 'r', encoding=encoding) as f:
                        return f.read()
                except:
                    pass

            logger.error("Failed to decode text file with multiple encodings")
            return ""
        except Exception as e:
            logger.error(f"Error reading text file: {e}")
            return ""

    def _process_text(self, text: str) -> Dict:
        """
        Process the extracted text to identify and structure resume information.

        Args:
            text: Extracted text from resume

        Returns:
            Dictionary containing structured resume data
        """
        result = {
            "contact_info": self._extract_contact_info(text),
            "skills": self._extract_skills(text),
            "experience": self._extract_experience(text),
            "education": self._extract_education(text),
            "summary": self._extract_summary(text),
            "projects": self._extract_projects(text),
            "certifications": self._extract_certifications(text),
            "full_text": text  # Store full text for additional processing
        }

        return result

    def _extract_contact_info(self, text: str) -> Dict:
        """Extract contact information from the text."""
        contact_info = {
            "name": "",
            "email": "",
            "phone": "",
            "linkedin": "",
            "github": "",
            "location": ""
        }

        # Basic name extraction (first 50 characters, assuming name is at the top)
        first_lines = text.split('\n')[:3]
        if first_lines:
            for line in first_lines:
                line = line.strip()
                if len(line) > 0 and len(line) < 50:  # Typical name length constraint
                    # Check if this line doesn't match other patterns
                    if not (self.email_pattern.search(line) or
                           self.phone_pattern.search(line) or
                           self.linkedin_pattern.search(line)):
                        contact_info["name"] = line
                        break

        # Extract email
        email_match = self.email_pattern.search(text)
        if email_match:
            contact_info["email"] = email_match.group(0)

        # Extract phone number
        phone_match = self.phone_pattern.search(text)
        if phone_match:
            contact_info["phone"] = phone_match.group(0)

        # Extract LinkedIn profile
        linkedin_match = self.linkedin_pattern.search(text)
        if linkedin_match:
            contact_info["linkedin"] = linkedin_match.group(0)

        # Extract GitHub profile
        github_match = self.github_pattern.search(text)
        if github_match:
            contact_info["github"] = github_match.group(0)

        # Extract location (more complex, using NLP)
        if self.nlp and SPACY_AVAILABLE:
            try:
                doc = self.nlp(text[:1000])  # Process first 1000 chars for efficiency
                for ent in doc.ents:
                    if ent.label_ == "GPE" and not contact_info["location"]:  # Geo-Political Entity
                        contact_info["location"] = ent.text
                        break
            except Exception as e:
                logger.warning(f"Error extracting location with NLP: {e}")

        return contact_info

    def _extract_skills(self, text: str) -> List[str]:
        """Extract skills from resume text."""
        # Identify the skills section (if present)
        skills = []

        # Try to find skills section
        text_lower = text.lower()
        skills_section = ""

        # Find potential skills section
        matches = list(self.skills_headers.finditer(text_lower))
        if matches:
            start_idx = matches[0].start()

            # Find the next section header to determine end of skills section
            next_section_match = None
            for pattern in [self.experience_headers, self.education_headers,
                           self.projects_headers, self.summary_headers]:
                for match in pattern.finditer(text_lower[start_idx + 10:]):  # +10 to skip current header
                    if next_section_match is None or match.start() < next_section_match.start():
                        next_section_match = match

            if next_section_match:
                skills_section = text_lower[start_idx:start_idx + 10 + next_section_match.start()]
            else:
                skills_section = text_lower[start_idx:]  # Take everything after skills header

        # If we couldn't identify a skills section, search the entire text
        if not skills_section:
            skills_section = text_lower

        # Match skills from our database
        for skill in self.skills_db:
            # Use word boundary to match whole words only
            skill_pattern = r'\b' + re.escape(skill) + r'\b'
            if re.search(skill_pattern, skills_section, re.IGNORECASE):
                skills.append(skill)

        # Also extract skill phrases using NLP
        if self.nlp and SPACY_AVAILABLE and len(skills) < 5:  # If we found few skills, try NLP approach
            try:
                doc = self.nlp(text)
                for chunk in doc.noun_chunks:
                    if (len(chunk.text.split()) <= 3 and  # Limit to 3 words
                        chunk.text.lower() not in [s.lower() for s in skills]):
                        # Check if this chunk might be a skill (not in stopwords, not a person, etc.)
                        if (not chunk.root.is_stop and
                            chunk.root.pos_ in ["NOUN", "PROPN"] and
                            not any(ent.text == chunk.text for ent in doc.ents if ent.label_ == "PERSON")):
                            skills.append(chunk.text.lower())
            except Exception as e:
                logger.warning(f"Error extracting skills with NLP: {e}")

        return sorted(list(set(skills)))  # Remove duplicates and sort

    def _extract_experience(self, text: str) -> List[Dict]:
        """Extract work experience from resume text."""
        experiences = []

        # Find the experience section
        text_lower = text.lower()
        matches = list(self.experience_headers.finditer(text_lower))

        if not matches:
            return experiences

        start_idx = matches[0].start()

        # Find the end of experience section (next major section)
        end_idx = len(text)
        for pattern in [self.education_headers, self.skills_headers,
                       self.projects_headers, self.certifications_headers]:
            matches = list(pattern.finditer(text_lower[start_idx:]))
            if matches and start_idx + matches[0].start() < end_idx:
                end_idx = start_idx + matches[0].start()

        experience_text = text[start_idx:end_idx]

        # Split experience into entries (approximation)
        # This could be improved with better parsing logic
        entries = re.split(r'\n\s*\n', experience_text)

        # Process each entry
        for entry in entries[1:]:  # Skip the header
            if len(entry.strip()) < 10:  # Skip very short entries
                continue

            experience = {
                "title": "",
                "company": "",
                "date_range": "",
                "description": entry.strip()
            }

            # Extract date ranges
            date_match = self.date_pattern.search(entry)
            if date_match:
                experience["date_range"] = date_match.group(0).strip()

            # Try to extract title and company
            lines = entry.strip().split('\n')
            if len(lines) >= 2:
                # Assume first line has title/company
                first_line = lines[0].strip()
                if "," in first_line:
                    parts = first_line.split(',', 1)
                    experience["title"] = parts[0].strip()
                    experience["company"] = parts[1].strip()
                elif " at " in first_line.lower():
                    parts = first_line.lower().split(" at ")
                    experience["title"] = first_line[:first_line.lower().find(" at ")].strip()
                    experience["company"] = first_line[first_line.lower().find(" at ")+4:].strip()
                else:
                    experience["title"] = first_line
                    if len(lines) >= 3:
                        experience["company"] = lines[1].strip()

            experiences.append(experience)

        return experiences

    def _extract_education(self, text: str) -> List[Dict]:
        """Extract education information from resume text."""
        education = []

        # Find the education section
        text_lower = text.lower()
        matches = list(self.education_headers.finditer(text_lower))

        if not matches:
            return education

        start_idx = matches[0].start()

        # Find the end of education section (next major section)
        end_idx = len(text)
        for pattern in [self.experience_headers, self.skills_headers,
                       self.projects_headers, self.certifications_headers]:
            matches = list(pattern.finditer(text_lower[start_idx:]))
            if matches and start_idx + matches[0].start() < end_idx:
                end_idx = start_idx + matches[0].start()

        education_text = text[start_idx:end_idx]

        # Split education into entries
        entries = re.split(r'\n\s*\n', education_text)

        # Process each entry
        for entry in entries[1:]:  # Skip the header
            if len(entry.strip()) < 10:  # Skip very short entries
                continue

            edu = {
                "degree": "",
                "institution": "",
                "date_range": "",
                "description": entry.strip()
            }

            # Extract date ranges
            date_match = self.date_pattern.search(entry)
            if date_match:
                edu["date_range"] = date_match.group(0).strip()

            # Try to extract degree and institution
            lines = entry.strip().split('\n')
            if len(lines) >= 1:
                # Look for common degree keywords
                degree_keywords = ["bachelor", "master", "phd", "doctorate", "bs", "ba", "ms", "ma", "mba"]
                for line in lines[:2]:  # Check first two lines
                    line_lower = line.lower()
                    if any(keyword in line_lower for keyword in degree_keywords):
                        edu["degree"] = line.strip()
                        # Try to get institution from next line if not in same line
                        if len(lines) > 1 and "university" in lines[1].lower() or "college" in lines[1].lower():
                            edu["institution"] = lines[1].strip()
                        break

                # If we didn't find a degree but found an institution
                if not edu["degree"] and not edu["institution"]:
                    for line in lines[:2]:
                        if "university" in line.lower() or "college" in line.lower():
                            edu["institution"] = line.strip()
                            break

            education.append(edu)

        return education

    def _extract_summary(self, text: str) -> str:
        """Extract professional summary or objective from resume."""
        summary = ""

        # Find summary section
        text_lower = text.lower()
        matches = list(self.summary_headers.finditer(text_lower))

        if not matches:
            return summary

        start_idx = matches[0].end()

        # Find the end of summary section (next major section)
        end_idx = len(text)
        for pattern in [self.experience_headers, self.education_headers,
                       self.skills_headers, self.projects_headers]:
            matches = list(pattern.finditer(text_lower[start_idx:]))
            if matches and start_idx + matches[0].start() < end_idx:
                end_idx = start_idx + matches[0].start()

        summary = text[start_idx:end_idx].strip()

        # Limit to first paragraph if it's too long
        if len(summary) > 500:
            para_end = summary.find('\n\n')
            if para_end > 0:
                summary = summary[:para_end].strip()

        return summary

    def _appears_to_be_image_pdf(self, extracted_text: str) -> bool:
        """
        Check if the extracted text appears to be from an image-based PDF.
        This is a heuristic check based on text quality.
        
        Args:
            extracted_text: Text extracted from PDF
            
        Returns:
            True if the PDF appears to be image-based
        """
        text = extracted_text.strip()
        
        # If no text was extracted, likely an image PDF
        if not text:
            return True
            
        # If text is too short, might be image PDF with poor OCR or just headers
        if len(text) < 100:
            return True
            
        # Check for common OCR artifacts or gibberish
        weird_char_count = len(re.findall(r'[^\w\s,.;:!?()[\]{}\'"-]', text))
        total_char_count = len(text)
        
        # If more than 15% of characters are weird, likely OCR artifacts
        if total_char_count > 0 and weird_char_count / total_char_count > 0.15:
            return True
            
        # Check for common resume sections - if none found, might be image PDF
        common_sections = ['experience', 'education', 'skills', 'profile', 'summary', 
                          'work', 'project', 'contact', 'qualification']
        section_matches = 0
        for section in common_sections:
            if re.search(r'\b' + re.escape(section) + r'\b', text.lower()):
                section_matches += 1
                
        # If less than 2 common resume sections found, likely an image PDF
        if section_matches < 2:
            return True
            
        return False
        
    def _is_image_based_pdf(self, pdf_path: str) -> bool:
        """
        Determine if a PDF is primarily image-based (scanned document).
        
        Args:
            pdf_path: Path to the PDF file
            
        Returns:
            True if the PDF appears to be image-based
        """
        try:
            # Use PyMuPDF (fitz) to analyze PDF structure
            doc = fitz.open(pdf_path)
            
            # Count pages with images vs pages with text
            image_page_count = 0
            total_pages = doc.page_count
            
            for page_num in range(total_pages):
                page = doc[page_num]
                
                # Check if page has text
                text = page.get_text()
                if len(text.strip()) < 50:  # If page has minimal text
                    # Check if page has images
                    image_list = page.get_images()
                    if len(image_list) > 0:
                        image_page_count += 1
                        
            # If more than 50% of pages are image-based, consider it an image PDF
            if total_pages > 0 and image_page_count / total_pages > 0.5:
                return True
                
            return False
        except Exception as e:
            logger.error(f"Error checking if PDF is image-based: {e}")
            # If we can't determine, assume it's not image-based
            return False
            
    def _extract_text_using_ocr(self, pdf_path: str) -> str:
        """
        Extract text from PDF using OCR for image-based PDFs.
        
        Args:
            pdf_path: Path to the PDF file
            
        Returns:
            Extracted text as string
        """
        if not self.use_ocr or not PYTESSERACT_AVAILABLE:
            logger.warning("OCR requested but pytesseract is not available")
            return ""
            
        try:
            logger.info("Starting OCR extraction from PDF")
            extracted_text = ""
            temp_dir = tempfile.mkdtemp()
            
            try:
                # Open PDF with PyMuPDF
                pdf_document = fitz.open(pdf_path)
                
                # Process each page
                for page_number in range(pdf_document.page_count):
                    # Get the page
                    page = pdf_document[page_number]
                    
                    # Convert page to image
                    pix = page.get_pixmap(dpi=300)  # Higher DPI for better OCR
                    image_path = os.path.join(temp_dir, f"page_{page_number}.png")
                    pix.save(image_path)
                    
                    # Perform OCR on the image
                    with Image.open(image_path) as img:
                        page_text = pytesseract.image_to_string(img, lang='eng')
                        extracted_text += page_text + "\n\n"
                        
                logger.success(f"Successfully extracted text from {pdf_document.page_count} pages using OCR")
                return extracted_text
                
            finally:
                # Clean up temporary directory
                shutil.rmtree(temp_dir, ignore_errors=True)
                
        except Exception as e:
            logger.error(f"Error during OCR extraction: {e}")
            return ""
            
    def _extract_projects(self, text: str) -> List[Dict]:
        """Extract projects from resume text."""
        projects = []

        # Find the projects section
        text_lower = text.lower()
        matches = list(self.projects_headers.finditer(text_lower))

        if not matches:
            return projects

        start_idx = matches[0].start()

        # Find the end of projects section (next major section)
        end_idx = len(text)
        for pattern in [self.experience_headers, self.education_headers,
                      self.skills_headers, self.certifications_headers]:
            matches = list(pattern.finditer(text_lower[start_idx:]))
            if matches and start_idx + matches[0].start() < end_idx:
                end_idx = start_idx + matches[0].start()

        projects_text = text[start_idx:end_idx]

        # Split projects into entries
        entries = re.split(r'\n\s*\n', projects_text)

        # Process each entry
        for entry in entries[1:]:  # Skip the header
            if len(entry.strip()) < 10:  # Skip very short entries
                continue

            project = {
                "name": "",
                "description": entry.strip(),
                "technologies": []
            }

            # Try to extract project name
            lines = entry.strip().split('\n')
            if lines:
                project["name"] = lines[0].strip()

            # Extract technologies mentioned (from skills list)
            for skill in self.skills_db:
                if re.search(r'\b' + re.escape(skill) + r'\b', entry, re.IGNORECASE):
                    project["technologies"].append(skill)

            projects.append(project)

        return projects

    def _extract_certifications(self, text: str) -> List[str]:
        """Extract certifications from resume text."""
        certifications = []

        # Define certification section headers
        self.certifications_headers = re.compile(r'certifications|certificates|qualifications', re.IGNORECASE)

        # Find the certifications section
        text_lower = text.lower()
        matches = list(self.certifications_headers.finditer(text_lower))

        if not matches:
            return certifications

        start_idx = matches[0].start()

        # Find the end of certifications section (next major section)
        end_idx = len(text)
        for pattern in [self.experience_headers, self.education_headers,
                      self.skills_headers, self.projects_headers]:
            matches = list(pattern.finditer(text_lower[start_idx:]))
            if matches and start_idx + matches[0].start() < end_idx:
                end_idx = start_idx + matches[0].start()

        cert_text = text[start_idx:end_idx]

        # Extract certification entries (often bullet points or lines)
        lines = cert_text.split('\n')
        for line in lines[1:]:  # Skip header
            line = line.strip()
            # Check for bullet points or numbering
            line = re.sub(r'^[\•\-\*\■\●\○\►\▶\⦿\⚫\⚪\✓\✔\✕\✖\✗\✘]\s*', '', line)
            line = re.sub(r'^\d+[\.\)]\s*', '', line)

            if len(line) > 5 and len(line) < 100:  # Typical certification name length
                certifications.append(line)

        return certifications
